from .models import *
from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.http import HttpResponseRedirect
from django.views import View
from .forms import StudentForm, ContactForm
import docx
import nltk


class DocumentView(View):
    def document_list(request):
        url = '/apps/first/documents/add'
        all_document_obj = Document.objects.all()
        lst = []
        for i in all_document_obj:
            current = {
            'PK': i.pk,
            'file_name': i.file_name,
            'author': i.author,
            'language_ID': i.language_ID
            }
            lst.append(current)
        return render(request, 'first/documents/documents_list.html', {'context': lst, 'url': url})

    def document_view(request, detail_view_id = 0):
        url = '/apps/first/documents/' + str(detail_view_id)
        doc = Document.objects.get(pk=detail_view_id)
        current = {
            'Type': 'Документ',
            'File Name': doc.file_name,
            'Author': doc.author,
            'Language_ID': doc.language_ID,
            }
        if request.POST.get('action') == 'Delete':
            doc.delete()
            return redirect('/apps/first/documents')
        elif request.POST.get('action') == 'Edit':
            return redirect(url + '/change')

        return render(request, 'first/documents/document_view_ajax.html', {'context': current, 'url': url})

    def document_add(request):
        url = '/apps/first/documents/add'
        
        if request.POST.get('action') == 'Save':
            doc = Document(
                file_name = request.POST.get('file_name'),
                author = request.POST.get('author'),
                language_ID = request.POST.get('language_id')
                )
            doc.save()
            return redirect('/apps/first/documents')
        return render(request, 'first/documents/document_add_ajax.html', {'context': url})

    def document_change(request, detail_view_id):
        url = '/apps/first/documents/' + str(detail_view_id)
        doc = Document.objects.get(pk=detail_view_id)

        if request.POST.get('action') == 'Delete':
            doc.delete()
            return redirect('/apps/first/documents')

        if request.POST.get('action') == 'Save':
            doc.file_name = request.POST.get('file_name')
            doc.author = request.POST.get('author')
            doc.language_ID = request.POST.get('language_id')
            doc.save()
            print("DONE!")
            return redirect('/apps/first/documents')

        current = {
            'Type': 'Документ',
            'file_name': doc.file_name,
            'author': doc.author,
            'language_id': doc.language_ID,
            }
        return render(request, 'first/documents/document_change_ajax.html', {'context': current, 'text': url})


class SentenceView(View):
    def sentence_list(request):
        all_sentence_obj = Sentence.objects.all()
        lst = []
        for i in all_sentence_obj:
            current = {
            'PK': i.pk,
            'ID': i.ID,
            'doc_position': i.doc_position,
            'text': i.text,
            #'document_ID': i.document_ID
            }
            lst.append(current)
        return render(request, 'first/sentences/sentences_list.html', {'context': lst})

    def sentence_view(request, detail_view_id):
        url = '/apps/first/sentences/' + str(detail_view_id)
        sentence = Sentence.objects.get(pk=detail_view_id)
        all_document_obj = Document.objects.all()

        for doc in all_document_obj:
            if doc.file_name == str(sentence.ID):
                print("DONE")
                doc_data = {
                'type': 'Документ',
                'file_name': doc.file_name,
                'author': doc.author,
                'language_id': doc.language_ID,
                }

        ##############################
        choice = Document.objects.all()
        arr = []
        for item in choice:
            var = {
                'id': item.pk,
                'value': item
            }
            arr.append(var)
        ###########################
        current = {
            'Type': 'Предложение',
            'ID': sentence.ID,
            'document_position': sentence.doc_position,
            'text': sentence.text,
            #Document ID ??
            }
        if request.POST.get('action') == 'Delete':
            sentence.delete()
            return redirect('/apps/first/sentences')
        elif request.POST.get('action') == 'Edit':
            return redirect(url + '/change')
        return render(request, 'first/sentences/sentence_view_ajax.html', {'context': current, 'choices': arr, 'url': url, 'doc_data': doc_data})
        #, 'doc_data': doc_data

    def sentence_add(request):
        url = "/apps/first/sentences"
        choice = Document.objects.all()
        arr = []
        print(request.POST)
        for item in choice:
            var = {
                'id': item.pk,
                'value': item
            }
            arr.append(var)
            
        number = 0
        if request.POST.get('doc_select'):
            var = request.POST.get('doc_select')
            number = var[7] + var[8]

        
        for doc in Document.objects.all():
            if doc.pk == int(number):
                sentence = Sentence(
                #Choice Menu
                ID = doc,
                doc_position = request.POST.get('doc_position'),
                text = request.POST.get('text'),
                document_ID = request.POST.get('document_id')
                )
                sentence.save()
                return redirect(url)
        return render(request, 'first/sentences/sentence_add_ajax.html', {'context': url, 'choices': arr})
    
    def sentence_change(request, detail_view_id):
        url = '/apps/first/documents/' + str(detail_view_id)
        sentence = Sentence.objects.get(pk=detail_view_id)
        #######
        choice = Document.objects.all()
        arr = []
        for item in choice:
            var = {
                'id': item.pk,
                'value': item
            }
            arr.append(var)
            
        number = 0
        if request.POST.get('document_id'):
            var = request.POST.get('document_id')
            number = var[7] + var[8]
        #######


        if request.POST.get('action') == 'Delete':
            sentence.delete()
            return redirect('/apps/first/sentences')
        

        if request.POST.get('action') == 'Save':
            for doc in Document.objects.all():
                if doc.pk == int(number):
                    sentence.ID = doc
                    sentence.doc_position = request.POST.get('file_name')
                    sentence.text = request.POST.get('author')
                    sentence.save()
            return redirect('/apps/first/sentences')

        obj = Sentence.objects.get(pk=detail_view_id)
        current = {
        'type': 'Предложение',
        'doc_position': obj.doc_position,
        'text': obj.text,
        'document_id': obj.document_ID
        }

        return render(request, 'first/sentences/sentence_change_ajax.html', {'context': current, 'choices': arr})


class ClauseView(View):
    def clause_list(request):
        clauses = Clause.objects.all()
        lst = []
        for i in clauses:
            current = {
            'PK': i.pk,
            'ID': i.ID,
            'Text': i.text,
            'Sentence_ID': i.sentence_ID,
            'Frame_ID': i.frame_ID
            }
            lst.append(current)
        return render(request, 'first/clauses/clauses.html', {'context': lst})

    def clause_add(request):
        url = "/apps/first/clauses"
        choice = Sentence.objects.all()
        arr = []
        print(request.POST)
        for item in choice:
            var = {
                'id': item.pk,
                'value': item
            }
            arr.append(var)
            
        number = 0
        if request.POST.get('doc_select'):
            var = request.POST.get('doc_select')
            number = var[7] + var[8]

        
        for sentence in Sentence.objects.all():
            if sentence.pk == int(number):
                print(request.POST)
                clause = Clause(
                #Choice Menu
                ID = sentence,
                text = request.POST.get('text'),
                sentence_ID = sentence,
                frame_ID = request.POST.get('frame_id')
                )
                clause.save()
                return redirect(url)
        return render(request, 'first/clauses/clause_add_ajax.html', {'context': url, 'choices': arr})

    def clause_change(request, detail_view_id):
        url = '/apps/first/clauses/' + str(detail_view_id)
        clause = Clause.objects.get(pk=detail_view_id)
        #######
        choice = Sentence.objects.all()
        arr = []
        for item in choice:
            var = {
                'id': item.pk,
                'value': item
            }
            arr.append(var)
            
        number = 0
        if request.POST.get('document_id'):
            var = request.POST.get('document_id')
            number = var[7] + var[8]
        #######


        if request.POST.get('action') == 'Delete':
            clause.delete()
            return redirect('/apps/first/clauses')
        

        if request.POST.get('action') == 'Save':
            for sentence in Sentence.objects.all():
                if sentence.pk == int(number):
                    print(request.POST)
                    clause = Clause(
                    #Choice Menu
                    ID = sentence,
                    text = request.POST.get('text'),
                    sentence_ID = sentence,
                    frame_ID = request.POST.get('frame_id')
                    )
                    clause.save()
                    return redirect(url)

        clause = Clause.objects.get(pk=detail_view_id)
        current = {
            'Type': 'Клауза',
            'sentence_id': clause.ID,
            'text': clause.text,
            'frame_id': clause.frame_ID,
            
            }

        return render(request, 'first/clauses/clause_change_ajax.html', {'context': current, 'choices': arr})

    def clause_view(request, detail_view_id):
        url = '/apps/first/clauses/' + str(detail_view_id)
        clause = Clause.objects.get(pk=detail_view_id)
        ##############################
        choice = Sentence.objects.all()
        arr = []
        for item in choice:
            var = {
                'id': item.pk,
                'value': item
            }
            arr.append(var)
        ###########################
        current = {
            'Type': 'Клауза',
            'sentence_id': clause.ID,
            'text': clause.text,
            'frame_id': clause.frame_ID,
            
            #Document ID ??
            }
        if request.POST.get('action') == 'Delete':
            clause.delete()
            return redirect('/apps/first/clauses')
        elif request.POST.get('action') == 'Edit':
            return redirect(url + '/change')
        return render(request, 'first/clauses/clause_view_ajax.html', {'context': current, 'choices': arr, 'url': url})


class SyntaxemeView(View):
    def syntaxeme_list(request):
        all_syntaxemes_obj = Syntaxeme.objects.all()
        lst = []
        for i in all_syntaxemes_obj:
            current = {
            'PK': i.pk,
            'ID': i.ID,
            'clause_ID': i.clause_ID,
            }
            lst.append(current)
        return render(request, 'first/syntaxemes/syntaxemes.html', {'context': lst})

    def syntaxeme_add(request):
        url = "/apps/first/syntaxemes"
        choice = Clause.objects.all()
        arr = []
        for item in choice:
            var = {
                'id': item.pk,
                'value': item
            }
            arr.append(var)
        ### 
        number = 0
        if request.POST.get('select'):
            var = request.POST.get('select')
            if var[8].isdigit():
                number = var[7] + var[8]
            else:
                number = var[7]
        ###   
        
        for clause in Clause.objects.all():
            if clause.pk == int(number):
                print(request.POST)
                syntaxeme = Syntaxeme(
                #Choice Menu
                ID = clause,
                clause_ID = clause,
                )
                syntaxeme.save()
                return redirect(url)
        return render(request, 'first/syntaxemes/syntaxeme_add_ajax.html', {'context': url, 'choices': arr})

    def syntaxeme_change(request, detail_view_id):
        url = '/apps/first/syntaxemes/' + str(detail_view_id)
        syntaxeme = Syntaxeme.objects.get(pk=detail_view_id)
        #######
        choice = Clause.objects.all()
        arr = []
        for item in choice:
            var = {
                'id': item.pk,
                'value': item
            }
            arr.append(var)
        ### 
        number = 0
        if request.POST.get('select'):
            var = request.POST.get('select')
            if var[8].isdigit():
                number = var[7] + var[8]
            else:
                number = var[7]
        ###   
        


        if request.POST.get('action') == 'Delete':
            syntaxeme.delete()
            return redirect('/apps/first/syntaxemes')
        

        if request.POST.get('action') == 'Save':
            for clause in Clause.objects.all():
                if clause.pk == int(number):
                    print(request.POST)
                    syntaxeme = Syntaxeme(
                    #Choice Menu
                    ID = clause,
                    clause_ID = clause,
                    )
                    syntaxeme.save()
                    return redirect(url)

        syntaxeme = Syntaxeme.objects.get(pk=detail_view_id)
        current = {
            'Type': 'Клауза',
            'clause_id': syntaxeme.ID,
            }

        return render(request, 'first/syntaxemes/syntaxeme_change_ajax.html', {'context': current, 'choices': arr})

    def syntaxeme_view(request, detail_view_id):
        url = '/apps/first/syntaxemes/' + str(detail_view_id)
        syntaxeme = Syntaxeme.objects.get(pk=detail_view_id)
        #######
        choice = Clause.objects.all()
        arr = []
        for item in choice:
            var = {
                'id': item.pk,
                'value': item
            }
            arr.append(var)
            
        number = 0
        if request.POST.get('select'):
            var = request.POST.get('select')
            number = var[7] + var[8]
        #######

        current = {
            'Type': 'Клауза',
            'clause_id': syntaxeme.ID,
            }
        if request.POST.get('action') == 'Delete':
            syntaxeme.delete()
            return redirect('/apps/first/syntaxemes')
        elif request.POST.get('action') == 'Edit':
            return redirect(url + '/change')
        return render(request, 'first/syntaxemes/syntaxeme_view_ajax.html', {'context': current, 'choices': arr, 'url': url})


class WordformView(View):
    def wordform_list(request):
        all_wordform_obj = Wordform.objects.all()
        lst = []
        for i in all_wordform_obj:
            current = {
            'PK': i.pk,
            'ID': i.ID,
            'syntaxeme_ID': i.syntaxeme_ID,
            'frame_role_ID': i.frame_role_ID,
            'position_in_document': i.position_in_document,
            'position_in_sentence': i.position_in_sentence,
            'clause_position': i.clause_position,
            'text': i.text,
            'root_morpheme_ID': i.root_morpheme_ID
            }
            lst.append(current)
        return render(request, 'first/wordforms/wordforms_list_ajax.html', {'context': lst})

    def wordform_add(request):
        url = "/apps/first/wordforms"
        choice = Syntaxeme.objects.all()
        arr = []
        for item in choice:
            var = {
                'id': item.pk,
                'value': item
            }
            arr.append(var)
        ### 
        number = 0
        if request.POST.get('select'):
            var = request.POST.get('select')
            if var[8].isdigit():
                number = var[7] + var[8]
            else:
                number = var[7]
        ###   
        print(request.POST)
        print(number)
        for syntaxeme in Syntaxeme.objects.all():
            if syntaxeme.pk == int(number):
                wordform = Wordform(
                #Choice Menu
                ID = syntaxeme,
                syntaxeme_ID = syntaxeme,
                frame_role_ID = request.POST.get('framerole_id'),
                position_in_document = request.POST.get('doc_position'),
                position_in_sentence = request.POST.get('sentence_position'),
                clause_position = request.POST.get('clause_position'),
                text = request.POST.get('text'),
                root_morpheme_ID = request.POST.get('root_morpheme_id'),
                )
                print("HELLLO")
                wordform.save()
                return redirect(url)
        return render(request, 'first/wordforms/wordforms_add_ajax.html', {'context': url, 'choices': arr})

    def wordform_change(request, detail_view_id):
        url = '/apps/first/syntaxemes/' + str(detail_view_id)
        syntaxeme = Syntaxeme.objects.get(pk=detail_view_id)
        #######
        choice = Clause.objects.all()
        arr = []
        for item in choice:
            var = {
                'id': item.pk,
                'value': item
            }
            arr.append(var)
        ### 
        number = 0
        if request.POST.get('select'):
            var = request.POST.get('select')
            if var[8].isdigit():
                number = var[7] + var[8]
            else:
                number = var[7]
        ###   
        


        if request.POST.get('action') == 'Delete':
            syntaxeme.delete()
            return redirect('/apps/first/syntaxemes')
        

        if request.POST.get('action') == 'Save':
            for clause in Clause.objects.all():
                if clause.pk == int(number):
                    print(request.POST)
                    syntaxeme = Syntaxeme(
                    #Choice Menu
                    ID = clause,
                    clause_ID = clause,
                    )
                    syntaxeme.save()
                    return redirect(url)

        syntaxeme = Syntaxeme.objects.get(pk=detail_view_id)
        current = {
            'Type': 'Клауза',
            'clause_id': syntaxeme.ID,
            }

        return render(request, 'first/wordforms/wordform_change_ajax.html', {'context': current, 'choices': arr})

    def wordform_view(request, detail_view_id):
        url = '/apps/first/syntaxemes/' + str(detail_view_id)
        syntaxeme = Syntaxeme.objects.get(pk=detail_view_id)
        #######
        choice = Clause.objects.all()
        arr = []
        for item in choice:
            var = {
                'id': item.pk,
                'value': item
            }
            arr.append(var)
            
        number = 0
        if request.POST.get('select'):
            var = request.POST.get('select')
            number = var[7] + var[8]
        #######

        current = {
            'Type': 'Клауза',
            'clause_id': syntaxeme.ID,
            }
        if request.POST.get('action') == 'Delete':
            syntaxeme.delete()
            return redirect('/apps/first/syntaxemes')
        elif request.POST.get('action') == 'Edit':
            return redirect(url + '/change')
        return render(request, 'first/wordforms/wordform_view_ajax.html', {'context': current, 'choices': arr, 'url': url})


def foo(request):
    print(request)
    # if request.method == "POST":
    #     print("DOne")
    #     return redirect('/apps/first/documents/upload')
    # return redirect('/apps/first/documents/upload')
    return render(request, 'first/ajax.html')
    # if request.method == "POST":
    #     form = ContactForm(request.POST)
    #     if form.is_valid():
    #         name = form.cleaned_data['name']
    #         form.save()
    #         return JsonResponse({"name": name}, status=200)
    #     else:
    #         errors = form.errors.as_json()
    #         return JsonResponse({"errors": errors}, status=400)
    
    # request.is_ajax() is deprecated since django 3.1
    # is_ajax = request.headers.get('X-Requested-With') == 'XMLHttpRequest'

    # if is_ajax:
    #     if request.method == 'POST':
    #         data = json.load(request)
    #         todo = data.get('payload')
    #         Todo.objects.create(task=todo['task'], completed=todo['completed'])
    #         return JsonResponse({'status': 'Todo added!'})
    #     return JsonResponse({'status': 'Invalid request'}, status=400)
    # else:
    #     return render(request, 'first/ajax.html')

    


def handle_uploaded_file(f, request):  
    with open('first/files/'+f.name, 'wb+') as destination:  
        for chunk in f.chunks():  
            destination.write(chunk) 
        file = Document(file_name = f.name, author = request.POST['author'], language_ID = 'rus')
        file.save()
        doc = docx.Document(f"first/files/{f.name}")
        all_paras = doc.paragraphs
        sentences = []

        for par in all_paras:
            token = nltk.sent_tokenize(par.text)
            print(type(token))
            if token != " ":
                for item in token:
                    string = "".join(item)
                    string = string.replace(u'\xa0', u' ')
                    sentences.append(string)
            
        
        print(sentences)
        for item in sentences:
            if item != ",":
                s = Sentence(ID = file, doc_position = "example", text = item, document_ID = 1)
                s.save()



def upload(request):  
    if request.method == 'POST': 
        print(request.POST)
        student = StudentForm(request.POST, request.FILES)  
        if student.is_valid():  
            handle_uploaded_file(request.FILES['file'], request)  
            return redirect('/apps/first/documents') 
    else:  
        student = StudentForm()  
        return render(request,"first/other/upload.html",{'form':student})  
    
def document_upload(request):
    doc = docx.Document("first/files/try.docx")
    all_paras = doc.paragraphs
    print(all_paras[5].text)
    return HttpResponse("Done!")
    # with open('files/Горький.doc', 'rb') as f:
    #     source_stream = StringIO(f.read())
    # document = Document(source_stream)
    # # print(source_stream)
    # source_stream.close()

