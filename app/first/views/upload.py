from app_api.models import *
from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.http import HttpResponseRedirect
from django.views import View
from first.forms import StudentForm
import docx
import nltk

def handle_uploaded_file(f, request):  
    with open('first/files/'+f.name, 'wb+') as destination:  
        for chunk in f.chunks():  
            destination.write(chunk) 
        file = Document(file_name = f.name, author = request.POST['author'], language_ID = 'rus')
        file.save()
        doc = docx.Document(f"first/files/{f.name}")
        all_paras = doc.paragraphs
        sentences = []

        for par in all_paras:
            token = nltk.sent_tokenize(par.text)
            print(type(token))
            if token != " ":
                for item in token:
                    string = "".join(item)
                    string = string.replace(u'\xa0', u' ')
                    sentences.append(string)
            
        
        print(sentences)
        for item in sentences:
            if item != ",":
                s = Sentence(ID = file, doc_position = "example", text = item, document_ID = 1)
                s.save()



def upload(request):  
    if request.method == 'POST': 
        print(request.POST)
        student = StudentForm(request.POST, request.FILES)  
        if student.is_valid():  
            handle_uploaded_file(request.FILES['file'], request)  
            return redirect('/documents') 
    else:  
        student = StudentForm()  
        return render(request,"first/other/upload.html",{'form':student})  
    
def document_upload(request):
    doc = docx.Document("first/files/try.docx")
    all_paras = doc.paragraphs
    print(all_paras[5].text)
    return HttpResponse("Done!")
    # with open('files/Горький.doc', 'rb') as f:
    #     source_stream = StringIO(f.read())
    # document = Document(source_stream)
    # # print(source_stream)
    # source_stream.close()

